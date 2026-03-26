"""
Test script: Adobe PDF Services API — PDF to DOCX conversion.

Usage:
    python test_adobe_convert.py <pdf_file> [--client-id YOUR_CLIENT_ID] [--client-secret YOUR_CLIENT_SECRET]

Or set environment variables:
    export ADOBE_CLIENT_ID=your_client_id
    export ADOBE_CLIENT_SECRET=your_client_secret
    python test_adobe_convert.py <pdf_file>

The converted DOCX is saved next to the input PDF and also to /tmp/.
"""

from __future__ import annotations

import argparse
import io
import os
import sys
import time
from pathlib import Path

from adobe.pdfservices.operation.auth.service_principal_credentials import (
    ServicePrincipalCredentials,
)
from adobe.pdfservices.operation.exception.exceptions import (
    SdkException,
    ServiceApiException,
    ServiceUsageException,
)
from adobe.pdfservices.operation.io.cloud_asset import CloudAsset
from adobe.pdfservices.operation.io.stream_asset import StreamAsset
from adobe.pdfservices.operation.pdf_services import PDFServices
from adobe.pdfservices.operation.pdf_services_media_type import PDFServicesMediaType
from adobe.pdfservices.operation.pdfjobs.jobs.export_pdf_job import ExportPDFJob
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_params import (
    ExportPDFParams,
)
from adobe.pdfservices.operation.pdfjobs.params.export_pdf.export_pdf_target_format import (
    ExportPDFTargetFormat,
)
from adobe.pdfservices.operation.pdfjobs.jobs.ocr_pdf_job import OCRPDFJob
from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_params import OCRParams
from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_locale import (
    OCRSupportedLocale,
)
from adobe.pdfservices.operation.pdfjobs.params.ocr_pdf.ocr_supported_type import (
    OCRSupportedType,
)
from adobe.pdfservices.operation.pdfjobs.result.export_pdf_result import (
    ExportPDFResult,
)
from adobe.pdfservices.operation.pdfjobs.result.ocr_pdf_result import OCRPDFResult


def convert_pdf_to_docx(
    pdf_bytes: bytes,
    client_id: str,
    client_secret: str,
    ocr: bool = False,
    ocr_locale: str = "EN_US",
) -> bytes:
    """Convert PDF to DOCX using Adobe PDF Services API.

    Args:
        pdf_bytes: Raw PDF file bytes.
        client_id: Adobe API client ID.
        client_secret: Adobe API client secret.
        ocr: If True, run OCR on the PDF first (for scanned documents).
             This adds a text layer to the PDF before converting to DOCX.
        ocr_locale: OCR language locale (default: EN_US). Supported locales:
             EN_US, EN_GB, FR_FR, DE_DE, ES_ES, IT_IT, JA_JP, KO_KR,
             ZH_CN, PT_BR, NL_NL, and many more (38 total).

    Returns:
        DOCX file bytes.
    """
    # 1. Authenticate
    credentials = ServicePrincipalCredentials(
        client_id=client_id,
        client_secret=client_secret,
    )
    pdf_services = PDFServices(credentials=credentials)

    # 2. Upload the PDF
    input_asset = pdf_services.upload(
        input_stream=io.BytesIO(pdf_bytes),
        mime_type=PDFServicesMediaType.PDF,
    )

    # 2.5 (Optional) OCR — for scanned PDFs, adds a text layer first
    if ocr:
        locale = getattr(OCRSupportedLocale, ocr_locale, OCRSupportedLocale.EN_US)
        ocr_params = OCRParams(
            ocr_locale=locale,
            ocr_type=OCRSupportedType.SEARCHABLE_IMAGE,
        )
        ocr_job = OCRPDFJob(input_asset=input_asset, ocr_params=ocr_params)
        ocr_location = pdf_services.submit(ocr_job)
        ocr_response = pdf_services.get_job_result(ocr_location, OCRPDFResult)
        # Use the OCR'd PDF as input for the export step
        input_asset = ocr_response.get_result().get_asset()

    # 3. Create export job (PDF → DOCX)
    export_pdf_params = ExportPDFParams(
        target_format=ExportPDFTargetFormat.DOCX,
    )
    export_pdf_job = ExportPDFJob(
        input_asset=input_asset,
        export_pdf_params=export_pdf_params,
    )

    # 4. Submit and poll
    location = pdf_services.submit(export_pdf_job)
    pdf_services_response = pdf_services.get_job_result(
        location,
        ExportPDFResult,
    )

    # 5. Download result
    result_asset: CloudAsset = pdf_services_response.get_result().get_asset()
    stream_asset: StreamAsset = pdf_services.get_content(result_asset)

    # Get bytes from the stream
    output = stream_asset.get_input_stream()
    if isinstance(output, bytes):
        return output
    return output.read()


def analyze_docx(docx_bytes: bytes) -> dict:
    """Quick analysis of the converted DOCX."""
    import zipfile

    from docx import Document

    bio = io.BytesIO(docx_bytes)

    # VML check
    with zipfile.ZipFile(bio) as z:
        xml = z.read("word/document.xml").decode()
        vml = xml.count("v:shape") + xml.count("v:textbox") + xml.count("wps:wsp")

    bio.seek(0)
    doc = Document(bio)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    bold = sum(1 for p in doc.paragraphs for r in p.runs if r.bold)
    chars = sum(len(p.text) for p in doc.paragraphs)
    tables = len(doc.tables)

    return {
        "vml_boxes": vml,
        "paragraphs": len(paras),
        "chars": chars,
        "bold_runs": bold,
        "tables": tables,
        "first_5": [p[:80] for p in paras[:5]],
    }


def convert_single(pdf_path: Path, out_dir: Path, client_id: str, client_secret: str, ocr: bool = False, ocr_locale: str = "EN_US") -> None:
    """Convert a single PDF and print analysis."""
    pdf_bytes = pdf_path.read_bytes()
    print(f"\nInput: {pdf_path.name} ({len(pdf_bytes) / 1024:.0f} KB)")

    print("  Converting via Adobe PDF Services API...")
    start = time.time()
    try:
        docx_bytes = convert_pdf_to_docx(pdf_bytes, client_id, client_secret, ocr=ocr, ocr_locale=ocr_locale)
    except (ServiceApiException, ServiceUsageException, SdkException) as e:
        print(f"  ERROR: {e}")
        return

    elapsed = time.time() - start
    print(f"  Done in {elapsed:.1f}s ({len(docx_bytes) / 1024:.0f} KB)")

    out_path = out_dir / f"{pdf_path.stem}.docx"
    out_path.write_bytes(docx_bytes)
    print(f"  Saved: {out_path}")

    analysis = analyze_docx(docx_bytes)
    print(f"  VML: {analysis['vml_boxes']}  Paras: {analysis['paragraphs']}  Chars: {analysis['chars']}  Bold: {analysis['bold_runs']}  Tables: {analysis['tables']}")
    for i, p in enumerate(analysis["first_5"][:3]):
        print(f"    [{i}] {p}")


def main():
    CLIENT_ID = os.environ.get("ADOBE_CLIENT_ID", "")
    CLIENT_SECRET = os.environ.get("ADOBE_CLIENT_SECRET", "")

    parser = argparse.ArgumentParser(description="Convert PDF(s) to DOCX using Adobe API")
    parser.add_argument(
        "path",
        nargs="?",
        default="/Users/vibhav.agrawal/Desktop/pdf testing",
        help="PDF file or directory of PDFs (default: ~/Desktop/pdf testing)",
    )
    parser.add_argument("--client-id", default=CLIENT_ID)
    parser.add_argument("--client-secret", default=CLIENT_SECRET)
    parser.add_argument("--ocr", action="store_true", help="Run OCR first (for scanned PDFs)")
    parser.add_argument("--ocr-locale", default="EN_US", help="OCR locale (default: EN_US)")
    args = parser.parse_args()

    input_path = Path(args.path)
    out_dir = Path("/tmp/adobe_eval")
    out_dir.mkdir(exist_ok=True)

    if input_path.is_file():
        pdfs = [input_path]
    elif input_path.is_dir():
        pdfs = sorted(input_path.glob("*.pdf")) or sorted(input_path.rglob("*.pdf"))
    else:
        print(f"ERROR: {input_path} not found")
        sys.exit(1)

    print(f"Found {len(pdfs)} PDFs. Output: {out_dir}/")
    print(f"{'=' * 60}")

    for pdf in pdfs:
        convert_single(pdf, out_dir, args.client_id, args.client_secret, ocr=args.ocr, ocr_locale=args.ocr_locale)

    print(f"\n{'=' * 60}")
    print(f"All done. Open results:")
    print(f"  open {out_dir}/")


if __name__ == "__main__":
    main()
